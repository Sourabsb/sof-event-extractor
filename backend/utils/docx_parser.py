"""
DOCX Parser Module
Handles extraction of text content from Microsoft Word documents
"""

from docx import Document
from pathlib import Path
import logging
from typing import Optional

logger = logging.getLogger(__name__)

class DocxParser:
    """Microsoft Word document text extraction utility"""
    
    def __init__(self):
        """Initialize DOCX parser"""
        pass
    
    def extract_text(self, file_path: Path) -> Optional[str]:
        """
        Extract text content from DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content or None if extraction fails
        """
        try:
            # Open the document
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # Combine all text
            full_text = '\n'.join(text_content)
            
            if not full_text.strip():
                logger.warning("No text content found in DOCX")
                return None
            
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {str(e)}")
            return None
    
    def get_document_properties(self, file_path: Path) -> dict:
        """
        Get document properties and metadata
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing document properties
        """
        try:
            doc = Document(file_path)
            properties = doc.core_properties
            
            info = {
                "title": properties.title,
                "author": properties.author,
                "subject": properties.subject,
                "created": properties.created.isoformat() if properties.created else None,
                "modified": properties.modified.isoformat() if properties.modified else None,
                "last_modified_by": properties.last_modified_by,
                "category": properties.category,
                "comments": properties.comments,
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables)
            }
            
            return {k: v for k, v in info.items() if v is not None}
            
        except Exception as e:
            logger.error(f"Failed to get document properties: {str(e)}")
            return {"error": str(e)}
    
    def extract_structured_content(self, file_path: Path) -> dict:
        """
        Extract structured content including headings and formatting
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with structured content
        """
        try:
            doc = Document(file_path)
            
            structured_content = {
                "headings": [],
                "paragraphs": [],
                "tables": [],
                "lists": []
            }
            
            # Extract paragraphs with style information
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    para_info = {
                        "text": paragraph.text.strip(),
                        "style": paragraph.style.name if paragraph.style else "Normal"
                    }
                    
                    # Check if it's a heading
                    if "heading" in paragraph.style.name.lower():
                        structured_content["headings"].append(para_info)
                    else:
                        structured_content["paragraphs"].append(para_info)
            
            # Extract tables with structure
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    "table_index": table_idx,
                    "rows": []
                }
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data["rows"].append(row_data)
                
                structured_content["tables"].append(table_data)
            
            return structured_content
            
        except Exception as e:
            logger.error(f"Failed to extract structured content: {str(e)}")
            return {"error": str(e)}
