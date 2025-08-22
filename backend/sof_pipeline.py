import io
import os
import re
import shutil
import time
from dataclasses import dataclass, field
from typing import List, Dict, Tuple
from datetime import datetime, timedelta

import json
import pandas as pd
import dateparser

# File parsing & OCR
import pdfplumber
import fitz  # PyMuPDF
from PIL import Image, ImageFilter, ImageOps
import pytesseract
from docx import Document

# Embeddings & LLM
from groq import Groq


# --------------------------
# Data structures
# --------------------------
@dataclass
class IngestedDoc:
    filename: str
    pages: List[str]
    combined_text: str


@dataclass
class LaytimeResult:
    events_df: pd.DataFrame
    laytime_allowed_days: float = 0.0
    laytime_consumed_days: float = 0.0
    laytime_saved_days: float = 0.0
    demurrage_due: float = 0.0
    dispatch_due: float = 0.0
    calculation_log: List[str] = field(default_factory=list)


# --------------------------
# Utilities
# --------------------------
def _ocr_image(img: Image.Image) -> str:
    if shutil.which("tesseract") is None:
        return ""
    try:
        gray = img.convert("L")
        w, h = gray.size
        if max(w, h) < 1800:
            scale = 2 if max(w, h) < 1000 else 1.5
            gray = gray.resize((int(w * scale), int(h * scale)))
        gray = ImageOps.autocontrast(gray)
        gray = gray.filter(ImageFilter.SHARPEN)
        bw = gray.point(lambda x: 0 if x < 155 else 255, mode="1")
        config = "--oem 1 --psm 6 -l eng"
        text = pytesseract.image_to_string(bw, config=config)
        if len(text.strip()) < 20:
            text = pytesseract.image_to_string(gray, config=config)
        return text
    except Exception:
        return ""


def _pdf_to_text_or_ocr(pdf_bytes: bytes) -> List[str]:
    pages: List[str] = []
    
    # First try: Extract text using pdfplumber
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():  # Only add non-empty pages
                    pages.append(text.strip())
        print(f"PDF text extraction successful. Extracted {len(pages)} pages with pdfplumber.")
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}")
        pages = []

    # Second try: If pdfplumber failed or extracted very little content, try PyMuPDF
    if not pages or sum(len(p.strip()) for p in pages) < 100:
        print("PDF appears to be scanned or pdfplumber extracted minimal content. Trying PyMuPDF...")
        pages = []
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for pno in range(doc.page_count):
                pg = doc.load_page(pno)
                
                # First try text extraction
                text = pg.get_text()
                if text.strip() and len(text.strip()) > 20:
                    pages.append(text.strip())
                    continue
                
                # If no text, try OCR
                print(f"No text found on page {pno + 1}, attempting OCR...")
                pix = pg.get_pixmap(dpi=300, alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = _ocr_image(img)
                if ocr_text.strip():
                    pages.append(ocr_text.strip())
                    
            print(f"PyMuPDF extraction completed. Extracted {len(pages)} pages.")
        except Exception as e:
            print(f"PyMuPDF extraction failed: {e}")
            
    return pages


def _docx_to_text(docx_bytes: bytes) -> str:
    try:
        f = io.BytesIO(docx_bytes)
        doc = Document(f)
        
        # Extract text from paragraphs
        paragraphs_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs_text.append(paragraph.text.strip())
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    tables_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = paragraphs_text + tables_text
        return "\n".join(all_text) if all_text else ""
        
    except Exception as e:
        print(f"Error extracting DOCX content: {e}")
        return ""


def _image_to_text(img_bytes: bytes) -> str:
    try:
        img = Image.open(io.BytesIO(img_bytes))
        return _ocr_image(img)
    except Exception:
        return ""


# --------------------------
# Public: file ingestion
# --------------------------
def process_uploaded_files(uploaded_files: List[object]) -> List[IngestedDoc]:
    docs: List[IngestedDoc] = []
    for f in uploaded_files:
        name = getattr(f, "name", "uploaded")
        ext = os.path.splitext(name)[1].lower()
        data = f.read() if hasattr(f, "read") else f.getvalue()

        print(f"Processing file: {name} (type: {ext}, size: {len(data)} bytes)")

        pages: List[str] = []
        if ext in [".pdf"]:
            pages = _pdf_to_text_or_ocr(data)
        elif ext in [".docx"]:
            try:
                docx_text = _docx_to_text(data)
                if docx_text.strip():  # Only add if we got actual content
                    pages = [docx_text]
                    print(f"DOCX text extracted: {len(docx_text)} characters")
                else:
                    print(f"No text extracted from DOCX file {name}")
                    pages = []
            except Exception as e:
                print(f"Error processing DOCX file {name}: {e}")
                pages = []
        elif ext in [".jpg", ".jpeg", ".png"]:
            image_text = _image_to_text(data)
            if image_text.strip():
                pages = [image_text]
                print(f"Image text extracted: {len(image_text)} characters")
            else:
                print(f"No text extracted from image {name}")
                pages = []
        elif ext in [".txt"]:
            try:
                text = data.decode("utf-8", errors="ignore")
                if text.strip():
                    pages = [text]
                    print(f"Text file processed: {len(text)} characters")
                else:
                    print(f"Empty text file: {name}")
                    pages = []
            except Exception as e:
                print(f"Error processing text file {name}: {e}")
                pages = []
        
        # Filter out None or empty strings before joining
        valid_pages = [p for p in pages if p and p.strip()]
        if not valid_pages:
            print(f"No valid content found in {name}")
            continue

        combined = "\n\n".join(valid_pages)
        print(f"Final combined text for {name}: {len(combined)} characters")
        docs.append(IngestedDoc(filename=name, pages=valid_pages, combined_text=combined))
    
    print(f"Total documents processed: {len(docs)}")
    return docs


# --------------------------
# Fallback extraction (simple regex)
# --------------------------
def _fallback_extract_events(text: str, filename: str) -> List[Dict]:
    events = []
    lines = text.splitlines()
    
    print(f"Starting fallback extraction for {filename} with {len(lines)} lines")
    
    # Enhanced time pattern matching
    time_patterns = [
        re.compile(r"\b(\d{1,2}):(\d{2})\b"),  # 14:30
        re.compile(r"\b(\d{4})\s*(?:hrs?|hours?|h)?\b", re.IGNORECASE),  # 1430 or 1430 hrs
        re.compile(r"\b(\d{1,2})\.(\d{2})\b"),  # 14.30
        re.compile(r"\b(\d{1,2})\s*:\s*(\d{2})\b"),  # 14 : 30 (with spaces)
    ]
    
    # Date pattern matching
    date_patterns = [
        re.compile(r"\b(\d{1,2})[\/\-\.](\d{1,2})[\/\-\.](\d{2,4})\b"),  # 21/05/2023, 21-05-2023, 21.05.2023
        re.compile(r"\b(\d{1,2})\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*(\d{2,4})?\b", re.IGNORECASE),
        re.compile(r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{1,2}),?\s*(\d{2,4})?\b", re.IGNORECASE),
        re.compile(r"\b(\d{1,2})(st|nd|rd|th)?\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*(\d{2,4})?\b", re.IGNORECASE),
    ]
    
    # Maritime event keywords to identify relevant lines
    maritime_keywords = [
        'pilot', 'berth', 'anchor', 'loading', 'discharge', 'cargo', 'hose', 'gangway',
        'commenced', 'completed', 'started', 'finished', 'connected', 'disconnected',
        'all fast', 'ready', 'notice', 'tender', 'survey', 'inspection', 'pump', 'tank',
        'manifest', 'customs', 'agent', 'master', 'chief', 'officer', 'weather', 'rain',
        'stop', 'delay', 'break', 'lunch', 'shift', 'crew', 'stevedore', 'terminal'
    ]
    
    current_date = None
    event_count = 0
    
    for line_num, line in enumerate(lines):
        line = line.strip()
        if not line or len(line) < 5:
            continue
            
        # Try to find date in current line
        for date_pattern in date_patterns:
            date_match = date_pattern.search(line)
            if date_match:
                try:
                    parsed_date = dateparser.parse(date_match.group())
                    if parsed_date:
                        current_date = parsed_date.strftime('%Y-%m-%d')
                        print(f"Found date {current_date} in line: {line[:50]}...")
                        break
                except:
                    continue
        
        # Look for time patterns
        found_time = False
        for time_pattern in time_patterns:
            time_matches = time_pattern.findall(line)
            if time_matches:
                for match in time_matches:
                    time_str = None
                    
                    if isinstance(match, tuple) and len(match) == 2:
                        hour, minute = match
                        try:
                            hour_int = int(hour)
                            minute_int = int(minute)
                            if 0 <= hour_int <= 23 and 0 <= minute_int <= 59:
                                time_str = f"{hour.zfill(2)}:{minute.zfill(2)}"
                        except:
                            continue
                    elif isinstance(match, str) and len(match) == 4 and match.isdigit():
                        # 1430 format
                        try:
                            hour_int = int(match[:2])
                            minute_int = int(match[2:])
                            if 0 <= hour_int <= 23 and 0 <= minute_int <= 59:
                                time_str = f"{match[:2]}:{match[2:]}"
                        except:
                            continue
                    
                    if time_str:
                        # Check if line contains maritime-related keywords
                        line_lower = line.lower()
                        is_maritime_event = any(keyword in line_lower for keyword in maritime_keywords)
                        
                        # Also include lines with time patterns that look like log entries
                        is_log_entry = bool(re.search(r'\b\d{1,2}[:/\.]\d{2}\b.*[a-zA-Z]', line))
                        
                        if is_maritime_event or is_log_entry or len(line) > 20:
                            events.append({
                                "filename": filename,
                                "event": line,  # Use the full line as event description
                                "start_time_iso": None,  # Will be processed later
                                "end_time_iso": None,
                                "laytime_counts": is_maritime_event and any(kw in line_lower for kw in ['loading', 'discharge', 'cargo', 'pump']),
                                "raw_line": line,
                                "extracted_time": time_str,
                                "extracted_date": current_date
                            })
                            event_count += 1
                            found_time = True
                            break
                
                if found_time:
                    break
    
    print(f"Fallback extraction completed for {filename}: found {event_count} events")
    return events


# --------------------------
# Event extraction (LLM-driven, format-agnostic)
# --------------------------
def _llm_extract_events_single(text: str, filename: str, groq_api_key: str, model: str = "llama-3.1-8b-instant") -> List[Dict]:
    try:
        client = Groq(api_key=groq_api_key)
        # Reduce snippet size to avoid rate limits
        snippet = text if len(text) <= 40000 else text[:40000]

        system_prompt = (
            "You are an expert maritime document analyzer specializing in Statement of Facts (SoF) extraction. "
            "Your task is to meticulously identify ALL events from maritime documents and return them as a valid JSON array. "
            "CRITICAL: Use EXACT event names as they appear in the document - do NOT paraphrase or summarize.\n\n"
            
            "EXTRACTION RULES:\n"
            "1. 'event': Use the EXACT wording from the document. Examples:\n"
            "   - If document says 'Pilot on board', use 'Pilot on board' (not 'Pilot boarded')\n"
            "   - If document says 'Commenced loading operations', use 'Commenced loading operations'\n"
            "   - If document says 'All fast', use 'All fast' (not 'Vessel secured')\n"
            "   - If document says 'Hose connected', use 'Hose connected'\n"
            "   - If document says 'Notice of Readiness tendered', use 'Notice of Readiness tendered'\n\n"
            
            "2. 'start_time': The start time in HH:MM format (24-hour). If only one time is mentioned, use it here.\n\n"
            
            "3. 'end_time': The end time in HH:MM format (24-hour). CRITICAL PAIRING RULES:\n"
            "   - For events like 'Commenced loading' at 14:30, actively look for 'Completed loading' later in the document\n"
            "   - When you find 'Completed loading' at 18:00, set the end_time of the 'Commenced loading' event to '18:00'\n"
            "   - For paired events, create TWO separate entries:\n"
            "     1) {\"event\": \"Commenced loading\", \"start_time\": \"14:30\", \"end_time\": \"18:00\", ...}\n"
            "     2) {\"event\": \"Completed loading\", \"start_time\": \"18:00\", \"end_time\": null, ...}\n"
            "   - Common pairs: Commenced→Completed, Started→Finished, Connected→Disconnected, Berthed→Unberthed\n"
            "   - If no corresponding end event exists, set end_time to null\n\n"
            
            "4. 'date': The date in YYYY-MM-DD format. Pay attention to date headers and context.\n\n"
            
            "5. 'laytime_counts': Boolean (true/false) indicating if this time counts against laytime:\n"
            "   - TRUE for: Loading, Discharging, Pumping, Cargo operations, Waiting for berth, Cargo handling\n"
            "   - FALSE for: Pilot operations, Weather delays, Equipment failures, Bunkering, Surveys, Documentation\n\n"
            
            "6. 'raw_line': The EXACT text line from which you extracted this event.\n\n"
            
            "COMPREHENSIVE EXTRACTION STRATEGY:\n"
            "- Read the ENTIRE document before extracting to identify event pairs\n"
            "- Extract EVERY timestamped event, no matter how minor\n"
            "- Look for patterns like 'From 14:30 to 18:00 loading operations' and split into start/end\n"
            "- Pay attention to different date formats (21/05, May 21, 21st May, etc.)\n"
            "- Look for time patterns: 0800, 08:00, 8.00, eight hundred hours\n"
            "- Include all operational events: pilot, berthing, cargo, delays, surveys, documentation\n\n"
            
            "OUTPUT FORMAT: Return ONLY a valid JSON array with no additional text.\n"
            "Example: [{\"event\": \"Commenced loading\", \"start_time\": \"14:30\", \"end_time\": \"18:00\", \"date\": \"2023-05-21\", \"laytime_counts\": true, \"raw_line\": \"21/05 1430 Commenced loading\"}, {\"event\": \"Completed loading\", \"start_time\": \"18:00\", \"end_time\": null, \"date\": \"2023-05-21\", \"laytime_counts\": true, \"raw_line\": \"21/05 1800 Completed loading\"}]"
        )
        user_prompt = (
            f"Extract all events from the following text from document '{filename}'. "
            "Pay close attention to dates mentioned anywhere in the text, as they provide context for subsequent events. "
            "CRITICALLY IMPORTANT: Look for event pairs (start/end) and link them properly with end_time values. "
            "Return ONLY a valid JSON array. Do not include any other text or explanations.\n\n"
            "Text:\n"
            f"```\n{snippet}\n```"
        )

        # Retry logic for rate limiting
        max_retries = 3
        retry_delay = 20  # seconds
        
        for attempt in range(max_retries):
            try:
                resp = client.chat.completions.create(
                    model=model,
                    temperature=0.0,
                    max_tokens=3000,  # Reduced to avoid rate limits
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    response_format={"type": "json_object"},
                )
                break  # Success, exit retry loop
            except Exception as e:
                if "rate_limit_exceeded" in str(e) and attempt < max_retries - 1:
                    print(f"Rate limit hit for {filename}. Waiting {retry_delay} seconds before retry {attempt + 1}...")
                    time.sleep(retry_delay)
                    continue
                else:
                    raise e  # Re-raise if not rate limit or last attempt

        content = resp.choices[0].message.content.strip()
        match = re.search(r"```json\s*([\s\S]*?)\s*```|(\[[\s\S]*?\]|\{[\s\S]*?\})", content)
        if not match:
            return []

        raw_json_text = match.group(1) or match.group(2)
        event_list = []

        try:
            data = json.loads(raw_json_text)
            if isinstance(data, dict) and "events" in data:
                event_list = data["events"]
            elif isinstance(data, list):
                event_list = data
        except json.JSONDecodeError as e:
            print(f"Warning: LLM produced malformed JSON for {filename} ({e}). Attempting to salvage individual events.")
            individual_objects = re.findall(r'\{[^{}]*\}', raw_json_text, re.DOTALL)
            for obj_str in individual_objects:
                try:
                    event_list.append(json.loads(obj_str.strip()))
                except json.JSONDecodeError:
                    continue

        if not event_list:
            return []

        normalized_events = []
        for item in event_list:
            if not isinstance(item, dict) or not item.get("event"):
                continue

            start_time_str = str(item.get("start_time") or "")
            end_time_str = str(item.get("end_time") or "")
            date_str = str(item.get("date") or "")

            start_iso = dateparser.parse(f"{date_str} {start_time_str}").isoformat() if date_str and start_time_str else None
            end_iso = dateparser.parse(f"{date_str} {end_time_str}").isoformat() if date_str and end_time_str else None

            if start_iso and end_iso and end_iso < start_iso:
                end_dt = pd.to_datetime(end_iso) + pd.Timedelta(days=1)
                end_iso = end_dt.isoformat()

            normalized_events.append({
                "filename": filename,
                "event": str(item.get("event")).strip(),
                "start_time_iso": start_iso,
                "end_time_iso": end_iso,
                "laytime_counts": item.get("laytime_counts", False),
                "raw_line": str(item.get("raw_line") or item.get("evidence", "")).strip(),
            })
        return normalized_events
    except Exception as e:
        print(f"Error during LLM extraction for {filename}: {e}")
        return []


def _llm_extract_summary(text: str, filename: str, groq_api_key: str, model: str = "llama-3.1-8b-instant") -> Dict[str, str]:
    """
    Uses a large language model to extract a summary of voyage details.
    """
    summary_data = {}
    try:
        client = Groq(api_key=groq_api_key)
        # Use a smaller snippet as summary info is usually at the top
        snippet = text if len(text) <= 10000 else text[:10000]

        system_prompt = (
            "You are an expert data extractor for maritime Statement of Facts (SoF) or laytime calculation documents. "
            "Your task is to identify key voyage information and return it as a valid JSON object. Do NOT invent information. "
            "If a value is not found, omit the key or set it to null.\n\n"
            "EXTRACT THESE FIELDS:\n"
            "- 'CREATED FOR' or 'VESSEL NAME': The name of the vessel (e.g., 'MV ALRAYAN').\n"
            "- 'VOYAGE FROM': The starting port or location of the voyage.\n"
            "- 'VOYAGE TO': The destination port or location.\n"
            "- 'CARGO': The type of cargo being transported.\n"
            "- 'PORT': The port where the current operations are happening.\n"
            "- 'OPERATION': The type of operation (e.g., 'Discharge', 'Loading').\n"
            "- 'DEMURRAGE': The demurrage rate, if specified (value only).\n"
            "- 'DISPATCH': The dispatch rate, if specified (value only).\n"
            "- 'LOAD/DISCH': The loading or discharging rate (e.g., '10000.00').\n"
            "- 'CARGO QTY': The total quantity of the cargo.\n\n"
            "CRITICAL: Return ONLY a single, valid JSON object. Do not include any other text, explanations, or markdown formatting."
        )
        user_prompt = (
            f"Extract the voyage summary from the following text from document '{filename}'.\n\n"
            "Text:\n"
            f"```\n{snippet}\n```"
        )

        resp = client.chat.completions.create(
            model=model,
            temperature=0.0,
            max_tokens=1024,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            response_format={"type": "json_object"},
        )

        content = resp.choices[0].message.content.strip()
        
        data = json.loads(content)
        if isinstance(data, dict):
            key_map = {
                'VESSEL NAME': 'CREATED FOR',
                'VESSEL': 'CREATED FOR',
                'LOAD/DISCH (MT/DAY)': 'LOAD/DISCH',
                'CARGO QTY (MT)': 'CARGO QTY'
            }
            
            normalized_data = {}
            for key, value in data.items():
                upper_key = key.upper()
                standard_key = key_map.get(upper_key, upper_key)
                normalized_data[standard_key] = str(value) if value is not None else ""
            
            summary_data = normalized_data

    except Exception as e:
        print(f"Error during LLM summary extraction for {filename}: {e}")
        return {}
        
    return summary_data


def _post_process_and_parse_times(events: List[Dict]) -> List[Dict]:
    processed_events = []
    date_context_map = {}

    for event in events:
        filename = event.get("filename")
        
        # If we already have start_time_iso, use it and update date context
        if event.get("start_time_iso"):
            try:
                dt = pd.to_datetime(event["start_time_iso"])
                date_context_map[filename] = dt.date()
            except (ValueError, TypeError):
                pass
            processed_events.append(event)
            continue

        # Handle fallback extracted time/date
        extracted_time = event.get("extracted_time")
        extracted_date = event.get("extracted_date")
        
        if extracted_time and extracted_date:
            try:
                iso_time = dateparser.parse(f"{extracted_date} {extracted_time}").isoformat()
                event["start_time_iso"] = iso_time
                dt = pd.to_datetime(iso_time)
                date_context_map[filename] = dt.date()
                processed_events.append(event)
                continue
            except:
                pass

        # Try to parse from raw_line with date context
        raw_line = event.get("raw_line", "")
        if not raw_line:
            processed_events.append(event)
            continue

        date_context = date_context_map.get(filename)
        iso_time = None
        try:
            settings = {'RELATIVE_BASE': datetime.combine(date_context, datetime.min.time())} if date_context else {}
            parsed_dt = dateparser.parse(raw_line, settings=settings, languages=['en'])
            if parsed_dt:
                if parsed_dt.year > 1900 and parsed_dt.month and parsed_dt.day:
                    date_context_map[filename] = parsed_dt.date()
                iso_time = parsed_dt.isoformat()
        except Exception:
            pass

        event["start_time_iso"] = iso_time
        processed_events.append(event)
        
    return processed_events


def _link_start_end_events(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty or 'start_time_iso' not in df.columns:
        return df

    df['_dt'] = pd.to_datetime(df['start_time_iso'], errors='coerce')
    df = df.sort_values(by=['filename', '_dt']).reset_index(drop=True)

    patterns = {
        'commenced': 'completed', 'started': 'finished', 'began': 'ended',
        'start': 'finish', 'begin': 'end',
    }
    all_keywords = set(patterns.keys()) | set(patterns.values())
    stop_words = {'of', 'the', 'a', 'an', 'at', 'on', 'for', 'to', 'by'}

    def get_task_from_event(event_str: str) -> str:
        words = event_str.lower().split()
        task_words = [word for word in words if word not in all_keywords and word not in stop_words]
        return " ".join(task_words)

    rows_to_drop = set()
    end_times = df['end_time_iso'].copy()

    for i, row in df.iterrows():
        if i in rows_to_drop or pd.notna(row['end_time_iso']):
            continue

        event_lower = row['event'].lower()
        for start_kw, end_kw in patterns.items():
            if start_kw in event_lower.split():
                start_task = get_task_from_event(event_lower)
                if not start_task:
                    continue

                for j, future_row in df.loc[i+1:].iterrows():
                    if j in rows_to_drop:
                        continue
                    if future_row['filename'] != row['filename']:
                        break
                    
                    future_event_lower = future_row['event'].lower()
                    if end_kw in future_event_lower.split():
                        end_task = get_task_from_event(future_event_lower)
                        if start_task == end_task:
                            end_times.loc[i] = future_row['start_time_iso']
                            rows_to_drop.add(j)
                            break
    
    df['end_time_iso'] = end_times
    if rows_to_drop:
        df = df.drop(index=list(rows_to_drop)).reset_index(drop=True)
    
    df = df.drop(columns=['_dt'], errors='ignore')
    return df


def _deduplicate_and_sort(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    
    df = df.dropna(subset=['start_time_iso'])
    df = df.sort_values(by=['filename', 'start_time_iso']).reset_index(drop=True)
    
    # Keep the last entry for any duplicate event descriptions and start times
    df = df.drop_duplicates(subset=['filename', 'event', 'start_time_iso'], keep='last')
    
    return df.reset_index(drop=True)


# --------------------------
# Public: Main extraction pipeline
# --------------------------
def extract_events_and_summary(docs: List[IngestedDoc], groq_api_key: str) -> Tuple[pd.DataFrame, Dict[str, str]]:
    all_events = []
    summary_data = {}
    
    if not groq_api_key:
        raise ValueError("Groq API key is not configured. Please set the GROQ_API_KEY environment variable.")

    for doc in docs:
        if not doc.combined_text.strip():
            print(f"Skipping document {doc.filename} because it is empty.")
            continue
            
        print(f"Processing document: {doc.filename} ({len(doc.combined_text)} characters)")
        
        # Step 1: Try LLM extraction first for events
        llm_events = _llm_extract_events_single(doc.combined_text, doc.filename, groq_api_key)
        
        # Step 2: If LLM fails or returns nothing, use fallback for events
        if not llm_events:
            print(f"LLM event extraction failed for {doc.filename}. Using enhanced fallback extraction.")
            fallback_events = _fallback_extract_events(doc.combined_text, doc.filename)
            if fallback_events:
                print(f"Fallback extraction found {len(fallback_events)} events for {doc.filename}")
            else:
                print(f"No events found even with fallback for {doc.filename}")
            all_events.extend(fallback_events)
        else:
            print(f"LLM extraction successful for {doc.filename}: {len(llm_events)} events found")
            all_events.extend(llm_events)

        # Step 3: Extract summary info, but only if we haven't found it yet.
        if not summary_data:
            try:
                summary_data = _llm_extract_summary(doc.combined_text, doc.filename, groq_api_key)
                if summary_data:
                    print(f"Summary extracted from {doc.filename}")
            except Exception as e:
                print(f"Summary extraction failed for {doc.filename}: {e}")
                summary_data = {}


    if not all_events:
        return pd.DataFrame(), summary_data

    # Step 4: Post-process to parse any missing timestamps
    processed_events = _post_process_and_parse_times(all_events)
    
    df = pd.DataFrame(processed_events)
    
    # Step 5: Link start/end events
    df = _link_start_end_events(df)
    
    # Step 6: Deduplicate and sort
    df = _deduplicate_and_sort(df)

    # Step 7: Final formatting before returning
    if not df.empty:
        # Convert to datetime and handle coercion
        df['start_time_iso'] = pd.to_datetime(df['start_time_iso'], errors='coerce')
        df['end_time_iso'] = pd.to_datetime(df['end_time_iso'], errors='coerce')

        # Add 'Day' column
        df['Day'] = df['start_time_iso'].dt.strftime('%a, %d %b').fillna('')

        # Define final column order
        final_columns = [
            'Day', 'event', 'start_time_iso', 'end_time_iso', 
            'laytime_counts', 'raw_line', 'filename'
        ]
        # Add columns that might exist, even if not in the ideal order
        for col in df.columns:
            if col not in final_columns:
                final_columns.append(col)
        
        # Filter to existing columns and reorder
        df = df[[col for col in final_columns if col in df.columns]]

        # Drop the source_page column if it exists
        if 'source_page' in df.columns:
            df = df.drop(columns=['source_page'])

    return df, summary_data


def safe_float(value: any, default: float = 0.0) -> float:
    """Safely convert a value to a float, returning a default on failure."""
    if value is None:
        return default
    try:
        # Handle strings with commas
        if isinstance(value, str):
            return float(value.replace(',', ''))
        return float(value)
    except (ValueError, TypeError):
        return default


def calculate_laytime(summary: Dict[str, any], events_df: pd.DataFrame) -> LaytimeResult:
    log = []
    
    # 1. Get key figures from summary, with robust conversion
    try:
        cargo_qty = safe_float(summary.get('CARGO QTY', 0) or summary.get('CARGO QTY (MT)', 0))
        load_disch_rate = safe_float(summary.get('LOAD/DISCH', 0) or summary.get('LOAD/DISCH (MT/DAY)', 0))
        demurrage_rate = safe_float(summary.get('DEMURRAGE', 0) or summary.get('DEMURRAGE ($/DAY)', 0))
        dispatch_rate = safe_float(summary.get('DISPATCH', 0))
        log.append(f"Inputs: Cargo Qty={cargo_qty}, Rate={load_disch_rate}, Demurrage={demurrage_rate}, Dispatch={dispatch_rate}")
    except (ValueError, TypeError) as e:
        return LaytimeResult(events_df=events_df, calculation_log=[f"Error: Invalid summary data. Please check inputs. Details: {e}"])

    # 2. Calculate Laytime Allowed
    if not load_disch_rate or load_disch_rate == 0:
        log.append("Error: Load/Discharge rate is zero, cannot calculate allowed laytime.")
        allowed_days = 0
    else:
        allowed_days = cargo_qty / load_disch_rate
        log.append(f"Laytime Allowed: {cargo_qty} MT / {load_disch_rate} MT/day = {allowed_days:.4f} days")

    # 3. Calculate Time Consumed from events
    if events_df.empty:
        log.append("Warning: No events found to calculate consumed time.")
        total_consumed_seconds = 0
    else:
        df = events_df.copy()
        df['start'] = pd.to_datetime(df['start_time_iso'], errors='coerce')
        df['end'] = pd.to_datetime(df['end_time_iso'], errors='coerce')
        
        df['duration'] = df.apply(lambda row: row['end'] - row['start'] if pd.notna(row['start']) and pd.notna(row['end']) else timedelta(0), axis=1)
        
        def format_duration(td):
            if pd.isna(td) or td.total_seconds() <= 0:
                return "00h 00m"
            days, remainder = divmod(td.total_seconds(), 86400)
            hours, remainder = divmod(remainder, 3600)
            minutes, _ = divmod(remainder, 60)
            if days > 0:
                return f"{int(days)}d {int(hours):02d}h {int(minutes):02d}m"
            return f"{int(hours):02d}h {int(minutes):02d}m"

        df['Duration'] = df['duration'].apply(format_duration)
        
        # Use the 'laytime_counts' boolean column for utilization
        df['laytime_utilization_%'] = df.apply(lambda row: 100 if row.get('laytime_counts', False) and row['duration'].total_seconds() > 0 else 0, axis=1)
        
        # Calculate consumed time only for events where laytime_counts is True
        df['consumed_duration'] = df.apply(lambda row: row['duration'] if row['laytime_utilization_%'] == 100 else timedelta(0), axis=1)
        
        valid_durations = df['consumed_duration'].dropna()
        total_consumed_seconds = sum(td.total_seconds() for td in valid_durations)
        
        events_df['Duration'] = df['Duration']
        events_df['Laytime Utilization %'] = df['laytime_utilization_%']

    consumed_days = total_consumed_seconds / (24 * 3600)
    log.append(f"Total Time Consumed: {consumed_days:.4f} days")

    # 4. Calculate Demurrage or Dispatch
    demurrage_due = 0
    dispatch_due = 0
    time_saved_days = 0
    
    if consumed_days > allowed_days:
        demurrage_days = consumed_days - allowed_days
        demurrage_due = demurrage_days * demurrage_rate
        log.append(f"Result: Demurrage. Days over: {demurrage_days:.4f}. Amount: ${demurrage_due:,.2f}")
    elif allowed_days > consumed_days:
        time_saved_days = allowed_days - consumed_days
        dispatch_due = time_saved_days * dispatch_rate
        log.append(f"Result: Dispatch. Days saved: {time_saved_days:.4f}. Amount: ${dispatch_due:,.2f}")
    else:
        log.append("Result: Vessel completed exactly on time.")

    # 5. Add running totals to DataFrame
    if not events_df.empty:
        df['consumed_seconds_cumulative'] = df['consumed_duration'].cumsum().dt.total_seconds()
        df['laytime_consumed'] = df['consumed_seconds_cumulative'] / (24 * 3600)
        df['laytime_remaining'] = allowed_days - df['laytime_consumed']
        
        events_df['Laytime Consumed (Days)'] = df['laytime_consumed'].map('{:,.4f}'.format)
        events_df['Laytime Remaining (Days)'] = df['laytime_remaining'].map('{:,.4f}'.format)

    result = LaytimeResult(
        events_df=events_df,
        laytime_allowed_days=allowed_days,
        laytime_consumed_days=consumed_days,
        laytime_saved_days=time_saved_days,
        demurrage_due=demurrage_due,
        dispatch_due=dispatch_due,
        calculation_log=log
    )
    
    return result